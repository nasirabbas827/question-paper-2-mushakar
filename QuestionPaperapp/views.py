from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Student, Course, StudyProgram

def login(request):
    if request.method == 'POST':
        email = request.POST['email']
        password = request.POST['password']
        try:
            student = Student.objects.get(email=email, password=password)
            request.session['student_id'] = student.id
            return redirect('student_dashboard')
        except Student.DoesNotExist:
            error_message = 'Invalid email or password.'
            return render(request, 'login.html', {'error_message': error_message})
    else:
        return render(request, 'login.html')

def student_dashboard(request):
    student_id = request.session.get('student_id')
    if student_id:
        student = Student.objects.get(id=student_id)
        courses = Course.objects.filter(study_program=student.study_program)
        return render(request, 'dashboard.html', {'student': student, 'courses': courses})
    else:
        return redirect('login')

def logout_user(request):
    request.session.flush()
    return redirect('login')

from django.shortcuts import render, redirect
from .forms import StudentSignupForm
from .models import StudyProgram

def student_signup(request):
    if request.method == 'POST':
        form = StudentSignupForm(request.POST)
        if form.is_valid():
            form.save()
            # Redirect to a success page or another URL
            return redirect('login')
    else:
        form = StudentSignupForm()
    
    # Fetch study programs from the database
    study_programs = StudyProgram.objects.all()
    
    return render(request, 'signup.html', {'form': form, 'study_programs': study_programs})
from django.shortcuts import render, redirect
from .models import Student, Course

def student_dashboard(request):
    student_id = request.session.get('student_id')
    if student_id:
        student = Student.objects.get(id=student_id)
        courses = Course.objects.filter(study_program=student.study_program)
        return render(request, 'dashboard.html', {'student': student, 'courses': courses})
    else:
        return redirect('login')

from django.shortcuts import render, redirect
from .models import Student, Course, Exam

def view_exams(request, course_id):
    student_id = request.session.get('student_id')
    if student_id:
        student = Student.objects.get(id=student_id)
        course = Course.objects.get(id=course_id)
        exams = Exam.objects.filter(study_program=course.study_program)
        return render(request, 'view_exams.html', {'student': student, 'course': course, 'exams': exams})
    else:
        return redirect('login')
    

from django.shortcuts import render, redirect
from .models import Student, Exam, Question

def view_difficulty_levels(request, exam_id):
    student_id = request.session.get('student_id')
    if student_id:
        student = Student.objects.get(id=student_id)
        exam = Exam.objects.get(id=exam_id)
        questions = Question.objects.filter(exam=exam)
        difficulty_levels = set(question.DifficultyLevel for question in questions)
        return render(request, 'difficulty_levels.html', {'student': student, 'exam': exam, 'difficulty_levels': difficulty_levels})
    else:
        return redirect('login')
    
from django.shortcuts import render, redirect
from .models import Student, Exam, Question
from random import choices

import io
from django.shortcuts import render, redirect
from .models import Student, Exam, Question
from docx import Document
from django.http import HttpResponse

def view_questions_by_difficulty(request, exam_id, difficulty):
    student_id = request.session.get('student_id')
    if student_id:
        student = Student.objects.get(id=student_id)
        exam = Exam.objects.get(id=exam_id)
        questions = Question.objects.filter(exam=exam, DifficultyLevel=difficulty)
        
        total_marks = 100

        # Create a list of question weights based on their weightage
        question_weights = [question.Weightage for question in questions]

        # Choose random questions based on their weightage until the total marks reach 100
        selected_questions = []
        while total_marks > 0:
            # Choose a question based on its weightage
            chosen_question = choices(questions, weights=question_weights, k=1)[0]

            # Add the chosen question to the list
            selected_questions.append(chosen_question)

            # Deduct the question's weightage from the total marks
            total_marks -= chosen_question.Weightage
        
        # Generate the DOCX file
        document = Document()
        document.add_heading(f'Questions for Difficulty Level "{difficulty}"', level=1)
        for idx, question in enumerate(selected_questions, start=1):
            document.add_paragraph(f'Question {idx}: {question.QuestionText}', style='BodyText')
            document.add_paragraph(f'Options:')
            document.add_paragraph(f'1. {question.Option1}')
            document.add_paragraph(f'2. {question.Option2}')
            document.add_paragraph(f'3. {question.Option3}')
            document.add_paragraph(f'4. {question.Option4}')
            document.add_paragraph('\n')

        # Save the document to a BytesIO object
        docx_output = io.BytesIO()
        document.save(docx_output)
        docx_output.seek(0)

        # Create a response object with the DOCX content
        response = HttpResponse(docx_output.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=questions_{difficulty}.docx'

        return response
    else:
        return redirect('login')

from django.shortcuts import render, redirect
from .models import PaperSubmission
from .forms import PaperSubmissionForm

def submit_paper(request, exam_id):
    student_id = request.session.get('student_id')
    if student_id:
        if request.method == 'POST':
            form = PaperSubmissionForm(request.POST, request.FILES)
            if form.is_valid():
                submission = form.save(commit=False)
                submission.student_id = student_id
                submission.exam_id = exam_id
                submission.save()
                return redirect('view_difficulty_levels', exam_id=exam_id)
        else:
            form = PaperSubmissionForm()
        return render(request, 'submit_paper.html', {'form': form})
    else:
        return redirect('login')  # Redirect to login if student is not authenticated

def view_marks(request):
    student_id = request.session.get('student_id')
    if student_id:
        submissions = PaperSubmission.objects.filter(student_id=student_id)
        return render(request, 'view_marks.html', {'submissions': submissions})
    else:
        return redirect('login')  
    
from django.shortcuts import render, redirect
from .models import Notification
from django.contrib.auth.models import User

def view_notifications(request):
    student_id = request.session.get('student_id')
    if student_id:
        notifications = Notification.objects.all()
        return render(request, 'view_notifications.html', {'notifications': notifications})
    else:
        return redirect('login')  # Redirect to login if student is not authenticated
